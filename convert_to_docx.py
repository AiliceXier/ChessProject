from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import re

def read_markdown(file_path):
    with open(file_path, 'r', encoding='utf-8') as f:
        return f.read()

def convert_md_to_docx(md_content, output_path):
    doc = Document()
    
    # 设置默认字体
    style = doc.styles['Normal']
    font = style.font
    font.name = '宋体'
    font.size = Pt(12)
    
    lines = md_content.split('\n')
    i = 0
    
    while i < len(lines):
        line = lines[i]
        
        # 跳过空行
        if not line.strip():
            i += 1
            continue
        
        # 处理标题
        if line.startswith('# '):
            # 一级标题
            title = line[2:].strip()
            p = doc.add_heading(title, level=1)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.font.name = '黑体'
                run.font.size = Pt(18)
                run.font.bold = True
        
        elif line.startswith('## '):
            # 二级标题
            title = line[3:].strip()
            p = doc.add_heading(title, level=2)
            for run in p.runs:
                run.font.name = '黑体'
                run.font.size = Pt(16)
                run.font.bold = True
        
        elif line.startswith('### '):
            # 三级标题
            title = line[4:].strip()
            p = doc.add_heading(title, level=3)
            for run in p.runs:
                run.font.name = '黑体'
                run.font.size = Pt(14)
                run.font.bold = True
        
        # 处理分隔线
        elif line.strip() == '***':
            doc.add_paragraph('_' * 50)
        
        # 处理表格
        elif line.startswith('|') and '---' not in line:
            # 收集表格所有行
            table_lines = []
            while i < len(lines) and lines[i].startswith('|'):
                if '---' not in lines[i]:
                    table_lines.append(lines[i])
                i += 1
            
            if table_lines:
                # 解析表格
                rows = []
                for table_line in table_lines:
                    cells = [cell.strip() for cell in table_line.split('|')[1:-1]]
                    rows.append(cells)
                
                if rows:
                    num_cols = len(rows[0])
                    table = doc.add_table(rows=len(rows), cols=num_cols)
                    table.style = 'Table Grid'
                    
                    for row_idx, row_data in enumerate(rows):
                        row = table.rows[row_idx]
                        for col_idx, cell_text in enumerate(row_data):
                            if col_idx < num_cols:
                                row.cells[col_idx].text = cell_text
            continue
        
        # 处理代码块
        elif line.startswith('```'):
            code_lines = []
            i += 1
            while i < len(lines) and not lines[i].startswith('```'):
                code_lines.append(lines[i])
                i += 1
            
            code_text = '\n'.join(code_lines)
            p = doc.add_paragraph()
            run = p.add_run(code_text)
            run.font.name = 'Courier New'
            run.font.size = Pt(10)
            
        # 处理列表
        elif line.strip().startswith(('1. ', '2. ', '3. ', '4. ', '5. ', '6. ', '7. ', '8. ', '9. ')):
            text = re.sub(r'^\d+\.\s*', '', line.strip())
            p = doc.add_paragraph(text, style='List Number')
        
        elif line.strip().startswith('- ') or line.strip().startswith('* '):
            text = line.strip()[2:]
            # 处理缩进
            indent_level = (len(line) - len(line.lstrip())) // 2
            p = doc.add_paragraph(text, style='List Bullet')
            if indent_level > 0:
                p.paragraph_format.left_indent = Inches(0.25 * indent_level)
        
        # 处理普通段落
        else:
            # 处理加粗 **text**
            text = line.strip()
            p = doc.add_paragraph()
            
            # 简单处理加粗
            parts = re.split(r'\*\*(.*?)\*\*', text)
            for idx, part in enumerate(parts):
                run = p.add_run(part)
                if idx % 2 == 1:  # 奇数索引是加粗内容
                    run.bold = True
        
        i += 1
    
    doc.save(output_path)
    print(f"文档已保存到: {output_path}")

if __name__ == '__main__':
    input_file = r'c:\Users\28399\Downloads\com.unity.services.samples.multiplayer-chess-cloud-code-main\com.unity.services.samples.multiplayer-chess-cloud-code-main\概要设计报告.md'
    output_file = r'd:\unity\chess\概要设计报告.docx'
    
    md_content = read_markdown(input_file)
    convert_md_to_docx(md_content, output_file)
