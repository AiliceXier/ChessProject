"""
Convert markdown report to .docx with proper formatting.
Handles: headings, tables, code blocks, blockquotes, lists, bold/italic, horizontal rules.
"""

import re, sys
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

def read_markdown(path):
    with open(path, 'r', encoding='utf-8') as f:
        return f.read()

def set_cell_shading(cell, color):
    """Set cell background color."""
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color}"/>')
    cell._tc.get_or_add_tcPr().append(shading)

def add_formatted_paragraph(doc, text, bold=False, font_name='宋体', font_size=12,
                            alignment=None, first_line_indent=None, space_after=Pt(6)):
    """Add a paragraph with given formatting."""
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.font.name = font_name
    run.font.size = Pt(font_size)
    r = run._element
    r.rPr.rFonts.set(qn('w:eastAsia'), font_name)
    if alignment:
        p.alignment = alignment
    if first_line_indent:
        p.paragraph_format.first_line_indent = first_line_indent
    p.paragraph_format.space_after = space_after
    return p

def render_inline_text(paragraph, text, default_size=12):
    """Parse inline markdown (**bold**, `code`) and add runs."""
    # Split by **bold** and `code` patterns
    pattern = r'(\*\*(.*?)\*\*|`(.*?)`)'
    last_end = 0
    for m in re.finditer(pattern, text):
        # Add text before this match
        if m.start() > last_end:
            run = paragraph.add_run(text[last_end:m.start()])
            run.font.size = Pt(default_size)
            run.font.name = '宋体'
            r = run._element
            r.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        if text[m.start():m.start()+2] == '**':
            # Bold
            run = paragraph.add_run(m.group(2))
            run.bold = True
            run.font.size = Pt(default_size)
            run.font.name = '宋体'
            r = run._element
            r.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        else:
            # Inline code
            run = paragraph.add_run(m.group(3))
            run.font.name = 'Courier New'
            run.font.size = Pt(default_size - 1)
        last_end = m.end()
    if last_end < len(text):
        run = paragraph.add_run(text[last_end:])
        run.font.size = Pt(default_size)
        run.font.name = '宋体'
        r = run._element
        r.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

def convert_md_to_docx(md_content, output_path):
    doc = Document()

    # Page setup
    section = doc.sections[0]
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.17)
    section.right_margin = Cm(3.17)

    # Default style
    style = doc.styles['Normal']
    font = style.font
    font.name = '宋体'
    font.size = Pt(12)
    style.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

    lines = md_content.split('\n')
    i = 0

    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        # Skip empty lines
        if not stripped:
            i += 1
            continue

        # Skip the --- separator lines
        if stripped == '---':
            i += 1
            continue

        # Horizontal rule
        if stripped == '***' or stripped == '___':
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(6)
            pPr = p._p.get_or_add_pPr()
            pBdr = parse_xml(
                f'<w:pBdr {nsdecls("w")}>'
                '  <w:bottom w:val="single" w:sz="6" w:space="1" w:color="999999"/>'
                '</w:pBdr>'
            )
            pPr.append(pBdr)
            i += 1
            continue

        # Heading 1
        if line.startswith('# '):
            title = line[2:].strip()
            p = doc.add_heading(title, level=1)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.font.name = '黑体'
                run.font.size = Pt(18)
                run.font.bold = True
                r = run._element
                r.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
            i += 1
            continue

        # Heading 2
        if line.startswith('## '):
            title = line[3:].strip()
            p = doc.add_heading(title, level=2)
            for run in p.runs:
                run.font.name = '黑体'
                run.font.size = Pt(16)
                run.font.bold = True
                r = run._element
                r.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
            i += 1
            continue

        # Heading 3
        if line.startswith('### '):
            title = line[4:].strip()
            p = doc.add_heading(title, level=3)
            for run in p.runs:
                run.font.name = '黑体'
                run.font.size = Pt(14)
                run.font.bold = True
                r = run._element
                r.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
            i += 1
            continue

        # Heading 4
        if line.startswith('#### '):
            title = line[5:].strip()
            p = doc.add_heading(title, level=4)
            for run in p.runs:
                run.font.name = '黑体'
                run.font.size = Pt(12)
                run.font.bold = True
                r = run._element
                r.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
            i += 1
            continue

        # Blockquote
        if stripped.startswith('>'):
            quote_texts = []
            while i < len(lines) and lines[i].strip().startswith('>'):
                qt = re.sub(r'^>\s?', '', lines[i].strip())
                quote_texts.append(qt)
                i += 1
            full_quote = '\n'.join(quote_texts)
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.5)
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(6)
            # Add shading to quote
            pPr = p._p.get_or_add_pPr()
            shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="F2F2F2" w:val="clear"/>')
            pPr.append(shd)
            render_inline_text(p, full_quote, 11)
            continue

        # Table (starts with | and has --- separator line)
        if stripped.startswith('|') and i + 1 < len(lines) and '---' in lines[i + 1]:
            table_lines = []
            # First line is header
            header_line = lines[i]
            i += 1
            # Skip separator line
            i += 1
            # Collect body lines
            while i < len(lines) and lines[i].strip().startswith('|'):
                table_lines.append(lines[i])
                i += 1

            # Parse header cells
            header_cells = [c.strip() for c in header_line.split('|')[1:-1]]
            num_cols = len(header_cells)

            # Parse body rows
            rows_data = []
            for tl in table_lines:
                cells = [c.strip() for c in tl.split('|')[1:-1]]
                if len(cells) == num_cols:
                    rows_data.append(cells)

            if num_cols > 0:
                table = doc.add_table(rows=1 + len(rows_data), cols=num_cols)
                table.alignment = WD_TABLE_ALIGNMENT.CENTER
                table.style = 'Table Grid'

                # Header row
                header_row = table.rows[0]
                for col_idx, cell_text in enumerate(header_cells):
                    cell = header_row.cells[col_idx]
                    # Clear default paragraph
                    cell.text = ''
                    p = cell.paragraphs[0]
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    run = p.add_run(cell_text)
                    run.bold = True
                    run.font.size = Pt(10)
                    run.font.name = '宋体'
                    r = run._element
                    r.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
                    # Shading
                    set_cell_shading(cell, "D9E2F3")

                # Body rows
                for row_idx, row_data in enumerate(rows_data):
                    row = table.rows[row_idx + 1]
                    for col_idx, cell_text in enumerate(row_data):
                        cell = row.cells[col_idx]
                        cell.text = ''
                        p = cell.paragraphs[0]
                        # Handle bold in cell
                        if cell_text.startswith('**') and cell_text.endswith('**'):
                            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                            run = p.add_run(cell_text[2:-2])
                            run.bold = True
                        else:
                            run = p.add_run(cell_text)
                        run.font.size = Pt(10)
                        run.font.name = '宋体'
                        r = run._element
                        r.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
                        # Alternate row shading
                        if row_idx % 2 == 1:
                            set_cell_shading(cell, "F2F2F2")

            # Add spacing after table
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(3)
            p.paragraph_format.space_after = Pt(3)
            continue

        # Code block (fenced)
        if stripped.startswith('```'):
            lang = stripped[3:].strip()
            code_lines = []
            i += 1
            while i < len(lines) and not lines[i].strip().startswith('```'):
                code_lines.append(lines[i])
                i += 1
            i += 1  # Skip closing ```

            if code_lines:
                code_text = '\n'.join(code_lines)
                p = doc.add_paragraph()
                p.paragraph_format.left_indent = Inches(0.3)
                p.paragraph_format.space_before = Pt(4)
                p.paragraph_format.space_after = Pt(4)
                # Add shading
                pPr = p._p.get_or_add_pPr()
                shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="F5F5F5" w:val="clear"/>')
                pPr.append(shd)
                run = p.add_run(code_text)
                run.font.name = 'Courier New'
                run.font.size = Pt(9)
            continue

        # Ordered list item (1. text)
        ordered_match = re.match(r'^(\s*)(\d+)\.\s+(.*)', line)
        if ordered_match:
            indent = len(ordered_match.group(1))
            num = ordered_match.group(2)
            text = ordered_match.group(3)
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.25 * (indent // 2 + 1))
            p.paragraph_format.space_after = Pt(2)
            run = p.add_run(f'{num}. ')
            render_inline_text(p, text, 12)
            i += 1
            continue

        # Unordered list item (- text or * text)
        unordered_match = re.match(r'^(\s*)[-*+]\s+(.*)', line)
        if unordered_match:
            indent = len(unordered_match.group(1))
            text = unordered_match.group(2)
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.25 * (indent // 2 + 1))
            p.paragraph_format.space_after = Pt(2)
            run = p.add_run('• ')
            run.font.name = '宋体'
            run.font.size = Pt(12)
            render_inline_text(p, text, 12)
            i += 1
            continue

        # [图X] placeholder
        fig_match = re.match(r'^\| +\[图(\d+)\]', stripped)
        if fig_match:
            fig_num = fig_match.group(1)
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_before = Pt(8)
            p.paragraph_format.space_after = Pt(8)
            # Add a shaded box as placeholder
            run = p.add_run(f'[ 图{fig_num}：此处插入截图 ]')
            run.font.size = Pt(10)
            run.font.color.rgb = RGBColor(150, 150, 150)
            run.italic = True
            # Add border
            pPr = p._p.get_or_add_pPr()
            pBdr = parse_xml(
                f'<w:pBdr {nsdecls("w")}>'
                '  <w:top w:val="dashed" w:sz="4" w:space="4" w:color="CCCCCC"/>'
                '  <w:bottom w:val="dashed" w:sz="4" w:space="4" w:color="CCCCCC"/>'
                '  <w:left w:val="dashed" w:sz="4" w:space="4" w:color="CCCCCC"/>'
                '  <w:right w:val="dashed" w:sz="4" w:space="4" w:color="CCCCCC"/>'
                '</w:pBdr>'
            )
            pPr.append(pBdr)
            # Skip the table row below
            i += 1
            # Skip the separator and next row
            while i < len(lines) and (not lines[i].strip() or lines[i].strip().startswith('|')):
                i += 1
            continue

        # Normal paragraph
        text = stripped
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(6)
        render_inline_text(p, text, 12)

        i += 1

    doc.save(output_path)
    print(f"Doc saved to: {output_path}")

if __name__ == '__main__':
    input_file = r'D:\unity\my_chess\AI使用实践报告.md'
    output_file = r'D:\unity\my_chess\课程设计报告.docx'

    md_content = read_markdown(input_file)
    convert_md_to_docx(md_content, output_file)
