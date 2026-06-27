"""
Professional docx converter v4:
- Embed PNG architecture diagrams
- Black table headers with white text
- Professional Chinese academic formatting
- Proper cover page
- Clean heading hierarchy
- Page breaks between major sections
"""

import re, os
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING, WD_BREAK
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

DIAGRAM_DIR = r'D:\unity\my_chess\diagrams\png'
DIAGRAMS = {
    '图3-1': os.path.join(DIAGRAM_DIR, '1_system_architecture.png'),
    '图3-2': os.path.join(DIAGRAM_DIR, '2_ai_routing.png'),
    '图3-3': os.path.join(DIAGRAM_DIR, '3_state_transition.png'),
    '图3-4': os.path.join(DIAGRAM_DIR, '4_online_dataflow.png'),
}

def set_run_font(run, font_name='宋体', font_size=12, bold=False, east_asia='宋体'):
    run.font.name = font_name
    run.font.size = Pt(font_size)
    run.bold = bold
    r = run._element
    rPr = r.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = parse_xml(f'<w:rFonts {nsdecls("w")} w:eastAsia="{east_asia}"/>')
        rPr.insert(0, rFonts)
    else:
        rFonts.set(qn('w:eastAsia'), east_asia)

def set_cell_shading(cell, color):
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color}"/>')
    cell._tc.get_or_add_tcPr().append(shading)

def set_para_spacing(p, line_spacing=1.5, before=0, after=3):
    pf = p.paragraph_format
    pf.line_spacing = line_spacing
    pf.space_before = Pt(before)
    pf.space_after = Pt(after)

def set_first_indent(p, indent=24):
    p.paragraph_format.first_line_indent = Pt(indent)

def add_page_break(doc):
    p = doc.add_paragraph()
    run = p.add_run()
    run.add_break(WD_BREAK.PAGE)

def render_inline(paragraph, text, font='宋体', size=12):
    pattern = r'(\*\*(.*?)\*\*|`(.*?)`)'
    last = 0
    for m in re.finditer(pattern, text):
        if m.start() > last:
            run = paragraph.add_run(text[last:m.start()])
            set_run_font(run, font, size)
        if text[m.start():m.start()+2] == '**':
            run = paragraph.add_run(m.group(2))
            set_run_font(run, font, size, bold=True)
        else:
            run = paragraph.add_run(m.group(3))
            set_run_font(run, 'Courier New', size-1)
        last = m.end()
    if last < len(text):
        run = paragraph.add_run(text[last:])
        set_run_font(run, font, size)

def add_body_para(doc, text, indent=True):
    p = doc.add_paragraph()
    set_para_spacing(p, 1.5, 0, 3)
    if indent:
        set_first_indent(p)
    render_inline(p, text, '宋体', 12)
    return p

def add_heading_styled(doc, text, level):
    """Add heading with professional styling."""
    if level == 1:
        add_page_break(doc)
        p = doc.add_heading(text, level=1)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in p.runs:
            set_run_font(run, '黑体', 18, bold=True, east_asia='黑体')
        set_para_spacing(p, 1.5, 12, 12)
    elif level == 2:
        p = doc.add_heading(text, level=2)
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        for run in p.runs:
            set_run_font(run, '黑体', 15, bold=True, east_asia='黑体')
        set_para_spacing(p, 1.5, 12, 6)
    elif level == 3:
        p = doc.add_heading(text, level=3)
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        for run in p.runs:
            set_run_font(run, '黑体', 13, bold=True, east_asia='黑体')
        set_para_spacing(p, 1.5, 8, 4)
    elif level == 4:
        p = doc.add_heading(text, level=4)
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        for run in p.runs:
            set_run_font(run, '黑体', 12, bold=True, east_asia='黑体')
        set_para_spacing(p, 1.5, 6, 3)
    return p

def add_table_from_rows(doc, headers, rows, col_widths=None):
    """Create a professional table with black header and alternating rows."""
    num_cols = len(headers)
    num_rows = len(rows)
    table = doc.add_table(rows=1 + num_rows, cols=num_cols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Table Grid'

    # Header row - TRANSPARENT background, black text
    for ci, h in enumerate(headers):
        cell = table.rows[0].cells[ci]
        cell.text = ''
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_para_spacing(p, 1.2, 2, 2)
        run = p.add_run(h)
        set_run_font(run, '宋体', 9, bold=True, east_asia='宋体')

    # Data rows
    for ri, row_data in enumerate(rows):
        for ci, cell_text in enumerate(row_data):
            cell = table.rows[ri + 1].cells[ci]
            cell.text = ''
            p = cell.paragraphs[0]
            set_para_spacing(p, 1.2, 1, 1)
            run = p.add_run(str(cell_text))
            set_run_font(run, '宋体', 9, east_asia='宋体')
            if ri % 2 == 1:
                set_cell_shading(cell, "F5F5F5")

    # Set column widths
    if col_widths:
        for row in table.rows:
            for ci, w in enumerate(col_widths):
                row.cells[ci].width = Cm(w)

    # Spacer
    sp = doc.add_paragraph()
    set_para_spacing(sp, 1.0, 2, 2)
    return table

def add_code_block(doc, code_text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.3)
    set_para_spacing(p, 1.0, 4, 4)
    pPr = p._p.get_or_add_pPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="F8F8F8" w:val="clear"/>')
    pPr.append(shd)
    run = p.add_run(code_text)
    set_run_font(run, 'Courier New', 8, east_asia='Courier New')

def add_diagram(doc, diag_key):
    """Embed a PNG diagram in the document."""
    if diag_key in DIAGRAMS and os.path.exists(DIAGRAMS[diag_key]):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_para_spacing(p, 1.0, 8, 8)
        run = p.add_run()
        run.add_picture(DIAGRAMS[diag_key], width=Inches(5.8))
        # Caption
        cap = doc.add_paragraph()
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_para_spacing(cap, 1.2, 2, 8)
        run = cap.add_run(f'{diag_key}  系统架构图')
        set_run_font(run, '宋体', 9, east_asia='宋体')


def convert(md_path, output_path):
    doc = Document()

    # Page setup
    sec = doc.sections[0]
    sec.top_margin = Cm(2.54)
    sec.bottom_margin = Cm(2.54)
    sec.left_margin = Cm(3.17)
    sec.right_margin = Cm(3.17)

    # Default style
    style = doc.styles['Normal']
    style.font.name = '宋体'
    style.font.size = Pt(12)
    style.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    style.paragraph_format.line_spacing = 1.5

    with open(md_path, 'r', encoding='utf-8') as f:
        lines = f.readlines()

    i = 0
    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        if not stripped:
            i += 1
            continue

        # Separator / HR
        if stripped in ('---', '***', '___'):
            i += 1
            continue

        # === COVER PAGE ===
        if stripped == '## 封面':
            i += 1
            cover_lines = []
            while i < len(lines) and not lines[i].startswith('#'):
                if lines[i].strip():
                    cover_lines.append(lines[i].strip())
                i += 1

            # School
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_before = Pt(60)
            run = p.add_run('信息与通信工程学院')
            set_run_font(run, '黑体', 22, bold=True, east_asia='黑体')

            # Course
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            set_para_spacing(p, 1.8, 6, 6)
            run = p.add_run('2024-2025-2  软件设计思想与方法II')
            set_run_font(run, '黑体', 16, east_asia='黑体')

            # Report title
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_before = Pt(36)
            p.paragraph_format.space_after = Pt(36)
            run = p.add_run('课程设计报告')
            set_run_font(run, '黑体', 28, bold=True, east_asia='黑体')

            # Info
            info_map = {}
            for cl in cover_lines:
                if '项目名称' in cl:
                    info_map['name'] = cl.split('：', 1)[-1].strip() if '：' in cl else cl.replace('项目名称', '').strip()
                elif '完成时间' in cl:
                    info_map['time'] = cl.split('：', 1)[-1].strip() if '：' in cl else cl.replace('完成时间', '').strip()
                elif '授课教师' in cl:
                    info_map['teacher'] = cl.split('：', 1)[-1].strip() if '：' in cl else cl.replace('授课教师', '').strip()

            for label, key in [('项目名称：', 'name'), ('完成时间：', 'time'), ('授课教师：', 'teacher')]:
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                set_para_spacing(p, 2.0, 6, 6)
                run = p.add_run(f'{label}{info_map.get(key, "")}')
                set_run_font(run, '宋体', 14, east_asia='宋体')

            # Members
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_before = Pt(24)
            run = p.add_run('小组成员：')
            set_run_font(run, '宋体', 14, east_asia='宋体')

            members = add_table_from_rows(doc,
                ['姓名', '学号', '分工', '贡献度'],
                [['（用户填写）', '（用户填写）', '（用户填写）', '（用户填写）'] for _ in range(3)],
                [3.5, 3.5, 4.5, 3.5])

            add_page_break(doc)
            continue

        # === TOC ===
        if stripped == '## 目录':
            i += 1
            toc_lines = []
            while i < len(lines) and not lines[i].startswith('#'):
                if lines[i].strip() and lines[i].strip() != '---':
                    toc_lines.append(lines[i].strip())
                i += 1

            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_after = Pt(24)
            run = p.add_run('目  录')
            set_run_font(run, '黑体', 18, bold=True, east_asia='黑体')

            for tl in toc_lines:
                indent_count = tl.count('&emsp;')
                text = tl.replace('&emsp;', '')
                p = doc.add_paragraph()
                p.paragraph_format.left_indent = Cm(indent_count * 0.8)
                set_para_spacing(p, 1.8, 2, 2)
                render_inline(p, text, '宋体', 12)

            add_page_break(doc)
            continue

        # ===== HEADINGS =====
        if line.startswith('##### '):
            add_heading_styled(doc, line[6:].strip(), 4)
            i += 1
            continue
        if line.startswith('#### '):
            add_heading_styled(doc, line[5:].strip(), 3)
            i += 1
            continue
        if line.startswith('### '):
            add_heading_styled(doc, line[4:].strip(), 2)
            i += 1
            continue
        if line.startswith('## '):
            title = line[3:].strip()
            if title not in ('封面', '目录'):
                add_heading_styled(doc, title, 1)
            i += 1
            continue
        if line.startswith('# '):
            i += 1
            continue

        # ===== BLOCKQUOTE =====
        if stripped.startswith('>'):
            qlines = []
            while i < len(lines) and lines[i].strip().startswith('>'):
                qlines.append(re.sub(r'^>\s?', '', lines[i].strip()))
                i += 1
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.4)
            set_para_spacing(p, 1.3, 3, 3)
            pPr = p._p.get_or_add_pPr()
            shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="F5F5F5" w:val="clear"/>')
            pPr.append(shd)
            render_inline(p, '\n'.join(qlines), '宋体', 11)
            continue

        # ===== TABLE =====
        if stripped.startswith('|') and i+1 < len(lines) and '---' in lines[i+1]:
            header_line = lines[i]
            i += 2  # skip header + separator
            body_lines = []
            while i < len(lines) and lines[i].strip().startswith('|'):
                body_lines.append(lines[i])
                i += 1

            headers = [c.strip() for c in header_line.split('|')[1:-1]]
            rows = []
            for bl in body_lines:
                cells = [c.strip() for c in bl.split('|')[1:-1]]
                if len(cells) == len(headers):
                    rows.append(cells)

            if rows:
                add_table_from_rows(doc, headers, rows)
            continue

        # ===== CODE BLOCK =====
        if stripped.startswith('```'):
            code_lines = []
            i += 1
            while i < len(lines) and not lines[i].strip().startswith('```'):
                code_lines.append(lines[i].rstrip())
                i += 1
            i += 1
            if code_lines:
                add_code_block(doc, '\n'.join(code_lines))
            continue

        # ===== DIAGRAM PLACEHOLDER =====
        fig_match = re.match(r'^\|\s*\[(图\d+-\d+)\]', stripped)
        if fig_match:
            diag_key = fig_match.group(1)
            add_diagram(doc, diag_key)
            i += 1
            while i < len(lines) and (not lines[i].strip() or lines[i].strip().startswith('|')):
                i += 1
            continue

        # ===== UNORDERED LIST =====
        ul_match = re.match(r'^(\s*)[-*+]\s+(.*)', stripped)
        if ul_match:
            indent = len(ul_match.group(1))
            text = ul_match.group(2)
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.25 * (indent//2 + 1))
            set_para_spacing(p, 1.5, 2, 2)
            run = p.add_run('• ')
            set_run_font(run, '宋体', 12)
            render_inline(p, text, '宋体', 12)
            i += 1
            continue

        # ===== ORDERED LIST =====
        ol_match = re.match(r'^(\s*)(\d+)\.\s+(.*)', stripped)
        if ol_match:
            indent = len(ol_match.group(1))
            num = ol_match.group(2)
            text = ol_match.group(3)
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.25 * (indent//2 + 1))
            set_para_spacing(p, 1.5, 2, 2)
            run = p.add_run(f'{num}. ')
            set_run_font(run, '宋体', 12)
            render_inline(p, text, '宋体', 12)
            i += 1
            continue

        # ===== BODY PARAGRAPH =====
        add_body_para(doc, stripped, indent=True)
        i += 1

    doc.save(output_path)
    print(f'Saved: {output_path}')


if __name__ == '__main__':
    convert(
        r'D:\unity\my_chess\full_report.md',
        r'D:\unity\my_chess\课程设计报告.docx'
    )
