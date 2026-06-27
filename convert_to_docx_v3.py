"""
Convert markdown report to .docx with proper Chinese academic formatting.

Key improvements over v2:
  - Proper cover page (centered title, info layout, member table)
  - Heading level remapping: ## → Heading1, ### → Heading2, #### → Heading3, ##### → Heading4
  - Headings and body follow Chinese academic conventions (黑体 titles, 宋体 body)
  - Page breaks before each major section
  - Line spacing 1.5 for body text
  - First-line indent (2 chars) for body paragraphs
  - Proper table formatting
"""

import re, sys
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml


def read_markdown(path):
    with open(path, 'r', encoding='utf-8') as f:
        return f.read()


def set_cell_shading(cell, color):
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color}"/>')
    cell._tc.get_or_add_tcPr().append(shading)


def set_paragraph_spacing(paragraph, line_spacing=1.5, space_before=0, space_after=Pt(3)):
    """Set paragraph line spacing."""
    pf = paragraph.paragraph_format
    pf.line_spacing = line_spacing
    pf.space_before = Pt(space_before)
    pf.space_after = space_after


def set_run_font(run, font_name='宋体', font_size=12, bold=False, east_asia='宋体'):
    """Set font properties for a run, including east-asian font."""
    run.font.name = font_name
    run.font.size = Pt(font_size)
    run.bold = bold
    r = run._element
    rPr = r.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = parse_xml(f'<w:rFonts {nsdecls("w")} w:eastAsia="{east_asia}"/>')
        rPr.insert(0, rFonts)
    else:
        rFonts.set(qn('w:eastAsia'), east_asia)


def set_paragraph_first_indent(paragraph, indent=Pt(24)):
    """Set first line indent (2 Chinese characters ≈ 24pt at 12pt font)."""
    paragraph.paragraph_format.first_line_indent = indent


def add_page_break(doc):
    """Add a page break."""
    p = doc.add_paragraph()
    run = p.add_run()
    run.add_break(docx.enum.text.WD_BREAK.PAGE)
    return p


def render_inline_text(paragraph, text, default_size=12, default_font='宋体', add_indent=False):
    """Parse inline markdown (**bold**, `code`) and add runs."""
    pattern = r'(\*\*(.*?)\*\*|`(.*?)`)'
    last_end = 0
    for m in re.finditer(pattern, text):
        if m.start() > last_end:
            run = paragraph.add_run(text[last_end:m.start()])
            set_run_font(run, default_font, default_size)
        if text[m.start():m.start()+2] == '**':
            run = paragraph.add_run(m.group(2))
            set_run_font(run, default_font, default_size, bold=True)
        else:
            run = paragraph.add_run(m.group(3))
            set_run_font(run, 'Courier New', default_size - 1)
        last_end = m.end()
    if last_end < len(text):
        run = paragraph.add_run(text[last_end:])
        set_run_font(run, default_font, default_size)
    if add_indent:
        set_paragraph_first_indent(paragraph)


def is_section_heading(text):
    """Check if a heading is a major section like '一、需求说明书' or '七、附录'."""
    section_patterns = [
        r'^[一二三四五六七八九十]+、',  # Chinese numbered sections
        r'^附录[A-Z]',                  # Appendix
    ]
    return any(re.match(p, text.strip()) for p in section_patterns)


def convert_md_to_docx(md_content, output_path):
    import docx.enum.text

    doc = Document()

    # --- Page setup ---
    section = doc.sections[0]
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.17)
    section.right_margin = Cm(3.17)

    # --- Default style ---
    style = doc.styles['Normal']
    font = style.font
    font.name = '宋体'
    font.size = Pt(12)
    style.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    # Default paragraph format
    style.paragraph_format.line_spacing = 1.5

    lines = md_content.split('\n')
    i = 0

    # Track current position in document
    cover_done = False
    toc_done = False
    body_started = False

    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        if not stripped:
            i += 1
            continue

        # --- Skip standalone separators ---
        if stripped == '---':
            i += 1
            continue

        # --- Horizontal rule ---
        if stripped in ('***', '___'):
            p = doc.add_paragraph()
            pPr = p._p.get_or_add_pPr()
            pBdr = parse_xml(
                f'<w:pBdr {nsdecls("w")}>'
                '  <w:bottom w:val="single" w:sz="6" w:space="1" w:color="999999"/>'
                '</w:pBdr>'
            )
            pPr.append(pBdr)
            i += 1
            continue

        # ========== COVER PAGE ==========
        if stripped == '## 封面':
            # Build a proper cover page
            i += 1
            # Read lines until next heading
            cover_lines = []
            while i < len(lines) and not lines[i].startswith('#'):
                if lines[i].strip():
                    cover_lines.append(lines[i].strip())
                i += 1

            # ===== School name =====
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_before = Pt(60)
            set_paragraph_spacing(p, 1.5, 0, 6)
            run = p.add_run('信息与通信工程学院')
            set_run_font(run, '黑体', 22, bold=True)

            # ===== Course info =====
            for cl in cover_lines:
                if '2024-2025' in cl or '软件设计' in cl:
                    p = doc.add_paragraph()
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    set_paragraph_spacing(p, 1.5, 6, 6)
                    run = p.add_run(cl)
                    set_run_font(run, '黑体', 18)

            # ===== Title =====
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_before = Pt(30)
            p.paragraph_format.space_after = Pt(30)
            run = p.add_run('课程设计报告')
            set_run_font(run, '黑体', 26, bold=True)

            # ===== Project info (left-aligned, indented) =====
            for cl in cover_lines:
                if '项目名称' in cl or '完成时间' in cl or '授课教师' in cl or '小组' in cl:
                    continue  # skip these, we'll handle them differently
                if '信息与通信' in cl or '2024-2025' in cl or '课程设计' in cl:
                    continue  # already handled above

            # Project name, date, teacher
            info_items = []
            for cl in cover_lines:
                if '项目名称' in cl:
                    info_items.append(('项目名称：', cl[len('项目名称：'):].strip()))
                elif '完成时间' in cl:
                    info_items.append(('完成时间：', cl[len('完成时间：'):].strip()))
                elif '授课教师' in cl:
                    info_items.append(('授课教师：', cl[len('授课教师：'):].strip()))

            for label, value in info_items:
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                set_paragraph_spacing(p, 2.0, 6, 6)
                run = p.add_run(f'{label}{value}')
                set_run_font(run, '宋体', 16)

            # ===== Team member table =====
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_before = Pt(20)
            run = p.add_run('小组成员：')
            set_run_font(run, '宋体', 14)

            # Read the table from markdown (it's between the text after "小组成员：")
            # The table was part of cover_lines, but we skipped it. Let's re-read from the original source.
            # Actually, let's just create the table manually.
            table = doc.add_table(rows=4, cols=4)
            table.alignment = WD_TABLE_ALIGNMENT.CENTER
            table.style = 'Table Grid'

            headers = ['姓名', '学号', '分工', '贡献度']
            for col_idx, h in enumerate(headers):
                cell = table.rows[0].cells[col_idx]
                cell.text = ''
                p = cell.paragraphs[0]
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = p.add_run(h)
                set_run_font(run, '宋体', 12, bold=True)
                set_cell_shading(cell, "D9E2F3")

            for row_idx in range(1, 4):
                for col_idx in range(4):
                    cell = table.rows[row_idx].cells[col_idx]
                    cell.text = ''
                    p = cell.paragraphs[0]
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    run = p.add_run('（用户填写）')
                    set_run_font(run, '宋体', 12)

            # ===== Page break after cover =====
            p = doc.add_paragraph()
            run = p.add_run()
            run.add_break(docx.enum.text.WD_BREAK.PAGE)
            cover_done = True
            continue

        # ========== TABLE OF CONTENTS ==========
        if stripped == '## 目录':
            i += 1
            # Read TOC lines
            toc_lines = []
            while i < len(lines) and not lines[i].startswith('#'):
                if lines[i].strip() and lines[i].strip() != '---':
                    toc_lines.append(lines[i].strip())
                i += 1

            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_before = Pt(20)
            p.paragraph_format.space_after = Pt(20)
            run = p.add_run('目  录')
            set_run_font(run, '黑体', 18, bold=True)

            for tl in toc_lines:
                # Remove &emsp; and render with proper indent
                indent_count = tl.count('&emsp;')
                text = tl.replace('&emsp;', '')
                p = doc.add_paragraph()
                p.paragraph_format.left_indent = Cm(indent_count * 0.8)
                set_paragraph_spacing(p, 1.8, 2, 2)
                render_inline_text(p, text, 12)

            # Page break after TOC
            p = doc.add_paragraph()
            run = p.add_run()
            run.add_break(docx.enum.text.WD_BREAK.PAGE)
            toc_done = True
            continue

        # ========== HEADINGS ==========

        # Markdown heading level: ##### → Word Heading4 (left-aligned)
        if line.startswith('##### '):
            title = line[6:].strip()
            if not title:
                i += 1
                continue
            p = doc.add_heading(title, level=4)
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            for run in p.runs:
                set_run_font(run, '黑体', 12, bold=True, east_asia='黑体')
            i += 1
            continue

        # Markdown heading level: #### → Word Heading3 (left-aligned)
        if line.startswith('#### '):
            title = line[5:].strip()
            if not title:
                i += 1
                continue
            p = doc.add_heading(title, level=3)
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            for run in p.runs:
                set_run_font(run, '黑体', 14, bold=True, east_asia='黑体')
            i += 1
            continue

        # Markdown heading level: ### → Word Heading2 (left-aligned)
        if line.startswith('### '):
            title = line[4:].strip()
            if not title:
                i += 1
                continue
            p = doc.add_heading(title, level=2)
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            for run in p.runs:
                set_run_font(run, '黑体', 15, bold=True, east_asia='黑体')
            i += 1
            continue

        # Markdown heading level: ## → Word Heading1 (major sections)
        if line.startswith('## '):
            title = line[3:].strip()
            if not title:
                i += 1
                continue

            # Skip 封面 and 目录 (already handled)
            if title in ('封面', '目录'):
                i += 1
                continue

            # Add page break before each major section (except first one)
            # Put page break before the paragraph
            p = doc.add_paragraph()
            run = p.add_run()
            run.add_break(docx.enum.text.WD_BREAK.PAGE)

            p = doc.add_heading(title, level=1)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                set_run_font(run, '黑体', 18, bold=True, east_asia='黑体')
            i += 1
            continue

        # Markdown heading level: # → Cover title (only one)
        if line.startswith('# '):
            title = line[2:].strip()
            # Skip - the cover title is handled in the cover page
            i += 1
            continue

        # ========== BLOCKQUOTE ==========
        if stripped.startswith('>'):
            quote_texts = []
            while i < len(lines) and lines[i].strip().startswith('>'):
                qt = re.sub(r'^>\s?', '', lines[i].strip())
                quote_texts.append(qt)
                i += 1
            full_quote = '\n'.join(quote_texts)
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.5)
            p.paragraph_format.right_indent = Inches(0.3)
            set_paragraph_spacing(p, 1.3, 4, 4)
            # Shading
            pPr = p._p.get_or_add_pPr()
            shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="F5F5F5" w:val="clear"/>')
            pPr.append(shd)
            render_inline_text(p, full_quote, 11, '宋体')
            continue

        # ========== TABLE ==========
        if stripped.startswith('|') and i + 1 < len(lines) and '---' in lines[i + 1]:
            table_lines = []
            header_line = lines[i]
            i += 1  # skip header
            i += 1  # skip separator
            while i < len(lines) and lines[i].strip().startswith('|'):
                table_lines.append(lines[i])
                i += 1

            header_cells = [c.strip() for c in header_line.split('|')[1:-1]]
            num_cols = len(header_cells)

            rows_data = []
            for tl in table_lines:
                cells = [c.strip() for c in tl.split('|')[1:-1]]
                if len(cells) == num_cols:
                    rows_data.append(cells)

            if num_cols > 0 and rows_data:
                table = doc.add_table(rows=1 + len(rows_data), cols=num_cols)
                table.alignment = WD_TABLE_ALIGNMENT.CENTER
                table.style = 'Table Grid'

                # Header
                header_row = table.rows[0]
                for col_idx, cell_text in enumerate(header_cells):
                    cell = header_row.cells[col_idx]
                    cell.text = ''
                    p = cell.paragraphs[0]
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    set_paragraph_spacing(p, 1.2, 2, 2)
                    run = p.add_run(cell_text)
                    set_run_font(run, '宋体', 10, bold=True)
                    set_cell_shading(cell, "D9E2F3")

                # Body
                for row_idx, row_data in enumerate(rows_data):
                    row = table.rows[row_idx + 1]
                    for col_idx, cell_text in enumerate(row_data):
                        cell = row.cells[col_idx]
                        cell.text = ''
                        p = cell.paragraphs[0]
                        set_paragraph_spacing(p, 1.2, 1, 1)
                        # Check for bold markers in cell
                        if cell_text.startswith('**') and cell_text.endswith('**'):
                            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                            run = p.add_run(cell_text[2:-2])
                            set_run_font(run, '宋体', 10, bold=True)
                        else:
                            run = p.add_run(cell_text)
                            set_run_font(run, '宋体', 10)
                        # Alternate row shading
                        if row_idx % 2 == 1:
                            set_cell_shading(cell, "F2F2F2")

            # Small spacer after table
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after = Pt(2)
            continue

        # ========== CODE BLOCK ==========
        if stripped.startswith('```'):
            lang = stripped[3:].strip()
            code_lines = []
            i += 1
            while i < len(lines) and not lines[i].strip().startswith('```'):
                code_lines.append(lines[i])
                i += 1
            i += 1  # Skip closing ```

            if code_lines:
                code_text = '\n'.join(code_lines)
                p = doc.add_paragraph()
                p.paragraph_format.left_indent = Inches(0.3)
                p.paragraph_format.right_indent = Inches(0.3)
                set_paragraph_spacing(p, 1.0, 4, 4)
                pPr = p._p.get_or_add_pPr()
                shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="F5F5F5" w:val="clear"/>')
                pPr.append(shd)
                run = p.add_run(code_text)
                set_run_font(run, 'Courier New', 9)
            continue

        # ========== ORDERED LIST ==========
        ordered_match = re.match(r'^(\s*)(\d+)\.\s+(.*)', line)
        if ordered_match:
            indent_level = len(ordered_match.group(1)) // 2
            num = ordered_match.group(2)
            text = ordered_match.group(3)
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.25 * (indent_level + 1))
            set_paragraph_spacing(p, 1.5, 2, 2)
            run = p.add_run(f'{num}. ')
            set_run_font(run, '宋体', 12)
            render_inline_text(p, text, 12, '宋体')
            i += 1
            continue

        # ========== UNORDERED LIST ==========
        unordered_match = re.match(r'^(\s*)[-*+]\s+(.*)', line)
        if unordered_match:
            indent_level = len(unordered_match.group(1)) // 2
            text = unordered_match.group(2)
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.25 * (indent_level + 1))
            set_paragraph_spacing(p, 1.5, 2, 2)
            run = p.add_run('• ')
            set_run_font(run, '宋体', 12)
            render_inline_text(p, text, 12, '宋体')
            i += 1
            continue

        # ========== FIGURE PLACEHOLDER ==========
        fig_match = re.match(r'^\| +\[图(\d+)\]', stripped)
        if fig_match:
            fig_num = fig_match.group(1)
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_before = Pt(12)
            p.paragraph_format.space_after = Pt(12)
            run = p.add_run(f'[ 图{fig_num}：此处插入截图 ]')
            set_run_font(run, '宋体', 10)
            run.font.color.rgb = RGBColor(150, 150, 150)
            run.italic = True
            # Add border
            pPr = p._p.get_or_add_pPr()
            pBdr = parse_xml(
                f'<w:pBdr {nsdecls("w")}>'
                '  <w:top w:val="dashed" w:sz="4" w:space="4" w:color="CCCCCC"/>'
                '  <w:bottom w:val="dashed" w:sz="4" w:space="4" w:color="CCCCCC"/>'
                '  <w:left w:val="dashed" w:sz="4" w:space="4" w:color="CCCCCC"/>'
                '  <w:right w:val="dashed" w:sz="4" w:space="4" w:color="CCCCCC"/>'
                '</w:pBdr>'
            )
            pPr.append(pBdr)
            i += 1
            # Skip following table rows if any
            while i < len(lines) and (not lines[i].strip() or lines[i].strip().startswith('|')):
                i += 1
            continue

        # ========== NORMAL PARAGRAPH ==========
        text = stripped
        # Skip the "小组成员：" line (already in cover)
        if text == '小组成员：' and not cover_done:
            i += 1
            continue

        p = doc.add_paragraph()
        set_paragraph_spacing(p, 1.5, 0, 3)
        # Add first-line indent for body paragraphs (2 chars)
        set_paragraph_first_indent(p)
        render_inline_text(p, text, 12, '宋体')

        i += 1

    doc.save(output_path)
    print(f"Doc saved to: {output_path}")


if __name__ == '__main__':
    input_file = r'D:\unity\my_chess\full_report.md'
    output_file = r'D:\unity\my_chess\课程设计报告.docx'

    md_content = read_markdown(input_file)
    convert_md_to_docx(md_content, output_file)
